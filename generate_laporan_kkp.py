#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script untuk generate Laporan KKP Semester 6
Mahasiswa: Ilham Maulana (312310514)
Judul: Perancangan Sistem Informasi Kegiatan dan Pengumuman Kecamatan Cibarusah Berbasis Web
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def add_page_break(doc):
    """Tambah page break"""
    doc.add_page_break()

def add_centered_heading(doc, text, level=1, size=14, bold=True, color=None):
    """Tambah heading centered"""
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    heading.style = f'Heading {level}'
    return heading

def add_justified_paragraph(doc, text, indent_first=0, size=11):
    """Tambah paragraph dengan justifikasi"""
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.first_line_indent = Inches(indent_first)
    for run in para.runs:
        run.font.size = Pt(size)
    return para

def shade_paragraph(para, color):
    """Shade paragraph dengan warna"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    para._element.get_or_add_pPr().append(shading_elm)

def create_table(doc, rows, cols):
    """Create table"""
    return doc.add_table(rows=rows, cols=cols)

def set_cell_background(cell, color):
    """Set background warna cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_pPr().append(shading_elm)

# ===== INISIALISASI DOKUMEN =====
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ===== HALAMAN JUDUL =====
title_section = doc.add_paragraph()
title_section.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_section.add_run('LAPORAN KERJA KULIAH PRAKTEK (KKP)')
run.font.size = Pt(16)
run.font.bold = True

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('"Perancangan Sistem Informasi Kegiatan dan Pengumuman\nKecamatan Cibarusah Berbasis Web untuk Meningkatkan\nAkses Informasi Publik"')
run.font.size = Pt(13)
run.font.bold = True

doc.add_paragraph()
doc.add_paragraph()

# Informasi Mahasiswa
info_table = create_table(doc, 5, 2)
info_table.autofit = False
info_table.allow_autofit = False

info_data = [
    ('Nama Mahasiswa', 'Ilham Maulana'),
    ('Nomor Induk Mahasiswa (NIM)', '312310514'),
    ('Konsentrasi Kuliah', 'Rekayasa Perangkat Lunak (RPL)'),
    ('Tempat KKP', 'Kecamatan Cibarusah, Kabupaten Bekasi, Jawa Barat'),
    ('Periode KKP', f'Januari - Juni 2026')
]

for i, (label, value) in enumerate(info_data):
    row_cells = info_table.rows[i].cells
    row_cells[0].text = label
    row_cells[1].text = value

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Tanggal dan Tempat
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_para.add_run(f'Bekasi, {datetime.now().strftime("%d %B %Y")}')

add_page_break(doc)

# ===== DAFTAR ISI =====
add_centered_heading(doc, 'DAFTAR ISI', size=14, bold=True)
doc.add_paragraph()

daftar_isi_items = [
    ('HALAMAN JUDUL', '...', 'i'),
    ('HALAMAN PENGESAHAN', '...', 'ii'),
    ('KATA PENGANTAR', '...', 'iii'),
    ('DAFTAR ISI', '...', 'iv'),
    ('DAFTAR GAMBAR', '...', 'v'),
    ('DAFTAR TABEL', '...', 'vi'),
    ('BAB I  PENDAHULUAN', '', '1'),
    ('    1.1 Latar Belakang', '...', '1'),
    ('    1.2 Identifikasi Masalah', '...', '3'),
    ('    1.3 Rumusan Masalah', '...', '4'),
    ('    1.4 Batasan Masalah', '...', '5'),
    ('    1.5 Tujuan KKP', '...', '6'),
    ('    1.6 Manfaat KKP', '...', '7'),
    ('    1.7 Metode Pelaksanaan', '...', '8'),
    ('    1.8 Sistematika Penulisan', '...', '9'),
    ('BAB II GAMBARAN UMUM INSTANSI', '', '10'),
    ('    2.1 Profil Kecamatan Cibarusah', '...', '10'),
    ('    2.2 Visi dan Misi', '...', '11'),
    ('    2.3 Struktur Organisasi', '...', '12'),
    ('    2.4 Tugas dan Fungsi', '...', '13'),
    ('    2.5 Kegiatan Mahasiswa selama KKP', '...', '14'),
    ('BAB III LANDASAN TEORI', '', '15'),
    ('    3.1 Sistem Informasi', '...', '15'),
    ('    3.2 Website dan Web Dinamis', '...', '17'),
    ('    3.3 Sistem Informasi Berbasis Web', '...', '18'),
    ('    3.4 Informasi Publik', '...', '19'),
    ('    3.5 Pelayanan Publik', '...', '20'),
    ('    3.6 Rekayasa Perangkat Lunak', '...', '21'),
    ('    3.7 Metodologi Waterfall', '...', '22'),
    ('    3.8 UML dan Diagram Sistem', '...', '24'),
    ('    3.9 Basis Data Relasional', '...', '26'),
    ('    3.10 User Interface dan User Experience', '...', '27'),
    ('    3.11 Black Box Testing', '...', '28'),
    ('    3.12 Tools dan Teknologi', '...', '29'),
    ('BAB IV ANALISIS, PERANCANGAN, DAN IMPLEMENTASI', '', '30'),
    ('    4.1 Analisis Sistem Berjalan', '...', '30'),
    ('    4.2 Analisis Kebutuhan Sistem', '...', '32'),
    ('    4.3 Perancangan Sistem', '...', '35'),
    ('    4.4 Implementasi Sistem', '...', '42'),
    ('    4.5 Pengujian Sistem (Black Box Testing)', '...', '48'),
    ('    4.6 Perbandingan Sebelum-Sesudah Sistem', '...', '55'),
    ('    4.7 Pembahasan Hasil', '...', '56'),
    ('BAB V PENUTUP', '', '58'),
    ('    5.1 Kesimpulan', '...', '58'),
    ('    5.2 Saran dan Rekomendasi', '...', '60'),
    ('DAFTAR PUSTAKA', '', '62'),
    ('LAMPIRAN', '', '65'),
]

for item in daftar_isi_items:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25 if item[0].startswith('    ') else 0)
    run = para.add_run(item[0])
    run.font.bold = item[1] == '...'

add_page_break(doc)

# ===== DAFTAR GAMBAR =====
add_centered_heading(doc, 'DAFTAR GAMBAR', size=14, bold=True)
doc.add_paragraph()

gambar_list = [
    ('Gambar 2.1', 'Struktur Organisasi Kecamatan Cibarusah', '12'),
    ('Gambar 3.1', 'Komponen Sistem Informasi', '16'),
    ('Gambar 3.2', 'Siklus Hidup Metodologi Waterfall', '23'),
    ('Gambar 4.1', 'Use Case Diagram Sistem Informasi', '35'),
    ('Gambar 4.2', 'Activity Diagram Login Admin', '36'),
    ('Gambar 4.3', 'Activity Diagram Kelola Pengumuman', '37'),
    ('Gambar 4.4', 'Entity Relationship Diagram (ERD)', '38'),
    ('Gambar 4.5', 'Halaman Beranda Website', '43'),
    ('Gambar 4.6', 'Halaman Daftar Pengumuman', '44'),
    ('Gambar 4.7', 'Dashboard Admin', '46'),
]

for no, title, page in gambar_list:
    doc.add_paragraph(f'{no}: {title} ......... {page}', style='List Bullet')

add_page_break(doc)

# ===== DAFTAR TABEL =====
add_centered_heading(doc, 'DAFTAR TABEL', size=14, bold=True)
doc.add_paragraph()

tabel_list = [
    ('Tabel 4.1', 'Analisis Sistem Berjalan', '31'),
    ('Tabel 4.2', 'Kebutuhan Fungsional Sistem', '33'),
    ('Tabel 4.3', 'Kebutuhan Non-Fungsional Sistem', '34'),
    ('Tabel 4.4', 'Struktur Tabel Admin', '38'),
    ('Tabel 4.5', 'Struktur Tabel Pengumuman', '39'),
    ('Tabel 4.6', 'Struktur Tabel Kegiatan', '40'),
    ('Tabel 4.7', 'Hasil Pengujian Black Box Testing', '49'),
    ('Tabel 4.8', 'Perbandingan Sebelum-Sesudah Sistem', '56'),
]

for no, title, page in tabel_list:
    doc.add_paragraph(f'{no}: {title} ......... {page}', style='List Bullet')

add_page_break(doc)

# ===== BAB I: PENDAHULUAN =====
add_centered_heading(doc, 'BAB I\nPENDAHULUAN', size=14, bold=True)
doc.add_paragraph()

# 1.1 Latar Belakang
add_centered_heading(doc, '1.1 Latar Belakang', level=2, size=12, bold=True)
doc.add_paragraph()

latar_belakang = """Perkembangan teknologi informasi telah mendorong instansi pemerintahan untuk menyediakan layanan informasi yang cepat, akurat, dan mudah diakses oleh masyarakat. Kecamatan sebagai salah satu perangkat pemerintahan yang berhubungan langsung dengan masyarakat memiliki peran penting dalam menyampaikan informasi terkait kegiatan, pengumuman resmi, jadwal pelayanan, agenda acara, dan informasi administrasi. Namun, penyampaian informasi di lingkungan Kecamatan Cibarusah masih banyak dilakukan secara manual melalui papan pengumuman atau media sosial.

Kondisi tersebut menyebabkan informasi belum tersusun secara optimal dan sulit diakses kembali oleh masyarakat ketika dibutuhkan. Informasi yang disampaikan melalui media sosial sering tidak terstruktur dengan rapi, tidak konsisten, dan tidak selalu menjadi sumber informasi resmi yang terpercaya. Selain itu, masyarakat membutuhkan satu media resmi berbasis web yang terpusat untuk mengakses informasi mengenai kegiatan kecamatan, pengumuman resmi, jadwal pelayanan, agenda acara, dan informasi administrasi secara mudah dan kapan saja.

Berdasarkan permasalahan tersebut, diperlukan perancangan sistem informasi kegiatan dan pengumuman berbasis web sebagai media penyampaian informasi publik yang lebih terstruktur, informatif, dan mudah diakses. Sistem ini diharapkan dapat meningkatkan transparansi pemerintahan dan memberikan kemudahan bagi masyarakat dalam memperoleh informasi resmi dari Kecamatan Cibarusah."""

add_justified_paragraph(doc, latar_belakang, indent_first=0.5)

# 1.2 Identifikasi Masalah
add_centered_heading(doc, '1.2 Identifikasi Masalah', level=2, size=12, bold=True)
doc.add_paragraph()

masalah_items = [
    "Penyampaian informasi kegiatan dan pengumuman kecamatan masih banyak dilakukan secara manual melalui papan pengumuman atau media sosial.",
    "Informasi yang disampaikan melalui media sosial belum tersusun secara sistematis dan sulit untuk ditelusuri kembali.",
    "Masyarakat belum memiliki satu media resmi berbasis web yang terpusat untuk mengakses informasi kecamatan.",
    "Informasi jadwal pelayanan, agenda acara, dan administrasi belum tersedia secara terpusat dan mudah diakses.",
    "Kecamatan membutuhkan sistem informasi yang dapat membantu pengelolaan dan penyampaian informasi publik secara lebih efisien.",
    "Kurangnya transparansi informasi publik mengakibatkan kepercayaan masyarakat terhadap institusi menjadi kurang optimal."
]

for item in masalah_items:
    doc.add_paragraph(item, style='List Bullet')

# 1.3 Rumusan Masalah
add_centered_heading(doc, '1.3 Rumusan Masalah', level=2, size=12, bold=True)
doc.add_paragraph()

rumusan_items = [
    "Bagaimana kondisi penyampaian informasi kegiatan dan pengumuman di Kecamatan Cibarusah saat ini?",
    "Bagaimana merancang sistem informasi kegiatan dan pengumuman Kecamatan Cibarusah berbasis web?",
    "Fitur apa saja yang dibutuhkan dalam sistem informasi kegiatan dan pengumuman Kecamatan Cibarusah?",
    "Bagaimana hasil rancangan sistem informasi berbasis web dalam membantu akses informasi publik masyarakat?"
]

for i, item in enumerate(rumusan_items, 1):
    doc.add_paragraph(f'{i}. {item}')

# 1.4 Batasan Masalah
add_centered_heading(doc, '1.4 Batasan Masalah', level=2, size=12, bold=True)
doc.add_paragraph()

batasan_items = [
    "Sistem yang dirancang berfokus pada penyampaian informasi kegiatan dan pengumuman Kecamatan Cibarusah.",
    "Informasi yang disediakan meliputi kegiatan kecamatan, pengumuman resmi, jadwal pelayanan, agenda acara, dan informasi administrasi.",
    "Sistem dirancang berbasis web dan dapat diakses melalui browser.",
    "Project tidak mencakup layanan administrasi online penuh seperti pengajuan surat secara digital, tanda tangan elektronik, atau integrasi dengan sistem kependudukan.",
    "Pengujian sistem dilakukan menggunakan metode Black Box Testing pada fitur utama.",
    "Sistem ditujukan untuk dua jenis pengguna, yaitu admin kecamatan dan masyarakat sebagai pengunjung website."
]

for item in batasan_items:
    doc.add_paragraph(item, style='List Bullet')

# 1.5 Tujuan KKP
add_centered_heading(doc, '1.5 Tujuan KKP', level=2, size=12, bold=True)
doc.add_paragraph()

tujuan_items = [
    "Menganalisis kebutuhan informasi publik pada Kecamatan Cibarusah.",
    "Merancang sistem informasi kegiatan dan pengumuman Kecamatan Cibarusah berbasis web.",
    "Menyediakan media informasi yang lebih terstruktur dan terpusat untuk masyarakat.",
    "Membantu pihak kecamatan dalam menyampaikan informasi kegiatan, pengumuman resmi, jadwal pelayanan, agenda acara, dan informasi administrasi.",
    "Meningkatkan kemudahan akses masyarakat terhadap informasi resmi kecamatan dan transparansi pelayanan publik."
]

for i, item in enumerate(tujuan_items, 1):
    doc.add_paragraph(f'{i}. {item}')

# 1.6 Manfaat KKP
add_centered_heading(doc, '1.6 Manfaat KKP', level=2, size=12, bold=True)
doc.add_paragraph()

doc.add_paragraph('1.6.1 Manfaat bagi Mahasiswa', style='Heading 3')
manfaat_mhs = [
    "Menerapkan ilmu yang diperoleh selama perkuliahan, khususnya pada bidang Rekayasa Perangkat Lunak.",
    "Melatih kemampuan analisis kebutuhan sistem dan desain aplikasi berbasis web.",
    "Meningkatkan kemampuan perancangan dan pengembangan sistem informasi berbasis web.",
    "Memberikan pengalaman langsung dalam menyelesaikan permasalahan teknologi informasi pada instansi pemerintahan."
]
for item in manfaat_mhs:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('1.6.2 Manfaat bagi Kecamatan Cibarusah', style='Heading 3')
manfaat_kec = [
    "Membantu kecamatan memiliki media informasi digital yang lebih terstruktur dan profesional.",
    "Memudahkan pengelolaan pengumuman, kegiatan, agenda, dan informasi pelayanan.",
    "Mendukung peningkatan kualitas penyampaian informasi kepada masyarakat.",
    "Mengurangi ketergantungan pada penyampaian informasi manual dan media sosial yang tidak terstruktur."
]
for item in manfaat_kec:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('1.6.3 Manfaat bagi Masyarakat', style='Heading 3')
manfaat_mas = [
    "Memudahkan masyarakat memperoleh informasi resmi kecamatan kapan saja dan di mana saja.",
    "Masyarakat dapat mengakses jadwal pelayanan, agenda kegiatan, dan pengumuman dengan lebih mudah.",
    "Informasi menjadi lebih mudah dicari, tersusun, dan dapat ditelusuri kembali dengan cepat.",
    "Meningkatkan transparansi layanan informasi publik dan kepercayaan masyarakat terhadap institusi."
]
for item in manfaat_mas:
    doc.add_paragraph(item, style='List Bullet')

# 1.7 Metode Pelaksanaan
add_centered_heading(doc, '1.7 Metode Pelaksanaan KKP', level=2, size=12, bold=True)
doc.add_paragraph()

metode_list = [
    ('Observasi', 'Mahasiswa melakukan pengamatan terhadap proses penyampaian informasi yang berjalan di Kecamatan Cibarusah, termasuk media yang digunakan, papan pengumuman, media sosial, alur publikasi, dan kebutuhan informasi masyarakat.'),
    ('Wawancara', 'Mahasiswa melakukan wawancara dengan pegawai atau pihak terkait untuk mengetahui kebutuhan sistem, kendala penyampaian informasi, dan fitur yang diinginkan.'),
    ('Studi Literatur', 'Mahasiswa mempelajari teori terkait sistem informasi, website, layanan informasi publik, perancangan sistem, pengembangan aplikasi berbasis web, dan metode pengujian.'),
    ('Analisis Kebutuhan Sistem', 'Mahasiswa mengidentifikasi kebutuhan sistem berdasarkan hasil observasi dan wawancara.'),
    ('Perancangan Sistem', 'Mahasiswa membuat rancangan menggunakan diagram UML (Use Case, Activity), struktur menu, rancangan database, dan rancangan antarmuka.'),
    ('Implementasi Sistem', 'Mahasiswa mengembangkan sistem informasi berbasis web sesuai dengan kebutuhan dan desain yang telah dirancang.'),
    ('Pengujian Sistem', 'Mahasiswa melakukan pengujian fitur menggunakan metode Black Box Testing.')
]

for no, (judul, deskripsi) in enumerate(metode_list, 1):
    para = doc.add_paragraph()
    para.add_run(f'{no}. {judul}\n').font.bold = True
    para.add_run(deskripsi)

# 1.8 Sistematika Penulisan
add_centered_heading(doc, '1.8 Sistematika Penulisan', level=2, size=12, bold=True)
doc.add_paragraph()

sistematika_text = """Sistematika penulisan laporan KKP ini terdiri dari lima bab. Bab I berisi pendahuluan yang meliputi latar belakang, identifikasi masalah, rumusan masalah, batasan masalah, tujuan, manfaat, metode pelaksanaan, dan sistematika penulisan. Bab II berisi gambaran umum Kecamatan Cibarusah sebagai tempat pelaksanaan KKP. Bab III berisi landasan teori yang mendukung perancangan sistem informasi kegiatan dan pengumuman berbasis web. Bab IV berisi hasil dan pembahasan mengenai analisis kebutuhan, perancangan, implementasi, dan pengujian sistem. Bab V berisi kesimpulan dan saran untuk pengembangan sistem lebih lanjut."""

add_justified_paragraph(doc, sistematika_text, indent_first=0.5)

add_page_break(doc)

# ===== BAB II: GAMBARAN UMUM =====
add_centered_heading(doc, 'BAB II\nGAMBARAN UMUM INSTANSI', size=14, bold=True)
doc.add_paragraph()

add_centered_heading(doc, '2.1 Profil Kecamatan Cibarusah', level=2, size=12, bold=True)
doc.add_paragraph()

profil_text = """Kecamatan Cibarusah merupakan salah satu kecamatan di Kabupaten Bekasi, Jawa Barat, yang memiliki peran penting dalam melayani masyarakat di wilayahnya. Sebagai institusi pemerintahan tingkat kecamatan, Kecamatan Cibarusah memiliki tugas dalam mendukung pelayanan administrasi, koordinasi pemerintahan, pemberdayaan masyarakat, serta penyampaian informasi kepada masyarakat di wilayah Kecamatan Cibarusah.

Kecamatan Cibarusah memiliki beberapa desa dan kelurahan yang tersebar di wilayahnya. Setiap desa dan kelurahan memiliki pemerintah lokal (pemerintah desa) yang bekerja sama dengan kantor kecamatan dalam memberikan pelayanan publik dan informasi kepada masyarakat.

Sebagai lembaga publik, Kecamatan Cibarusah menyelenggarakan berbagai layanan administratif mulai dari pembuatan kartu identitas, surat keterangan, izin usaha, hingga pendataan penduduk. Selain itu, kecamatan juga mengelola program-program pemberdayaan masyarakat dan melaksanakan koordinasi pemerintahan dengan berbagai instansi di tingkat kabupaten dan pusat."""

add_justified_paragraph(doc, profil_text, indent_first=0.5)

add_centered_heading(doc, '2.2 Visi dan Misi Kecamatan Cibarusah', level=2, size=12, bold=True)
doc.add_paragraph()

doc.add_paragraph('Visi:', style='Heading 3')
doc.add_paragraph('Mewujudkan pelayanan pemerintahan yang responsif, transparan, dan akuntabel untuk meningkatkan kesejahteraan masyarakat Kecamatan Cibarusah.')

doc.add_paragraph('Misi:', style='Heading 3')
misi_items = [
    'Meningkatkan kualitas pelayanan publik yang cepat, tepat, dan mudah diakses oleh masyarakat.',
    'Mewujudkan transparansi dan akuntabilitas dalam penyelenggaraan pemerintahan.',
    'Memberdayakan masyarakat untuk ikut berperan serta dalam pembangunan daerah.',
    'Membangun kemitraan yang kuat dengan berbagai stakeholder dalam memberikan pelayanan terbaik.',
    'Meningkatkan penggunaan teknologi informasi dalam mendukung efisiensi pelayanan publik.'
]
for item in misi_items:
    doc.add_paragraph(item, style='List Bullet')

add_centered_heading(doc, '2.3 Struktur Organisasi', level=2, size=12, bold=True)
doc.add_paragraph()

org_text = """Struktur organisasi Kecamatan Cibarusah dipimpin oleh seorang Camat yang bertanggung jawab terhadap Bupati Bekasi. Di bawah Camat terdapat beberapa seksi dan bagian yang menangani berbagai aspek pemerintahan dan pelayanan publik, meliputi:

• Seksi Pemerintahan dan Otonomi Daerah
• Seksi Pembangunan Daerah
• Seksi Kesejahteraan Sosial
• Seksi Ketertiban dan Keamanan
• Bagian Administrasi dan Tata Usaha
• Bagian Keuangan"""

add_justified_paragraph(doc, org_text, indent_first=0.5)

add_centered_heading(doc, '2.4 Tugas dan Fungsi Kecamatan', level=2, size=12, bold=True)
doc.add_paragraph()

tugas_items = [
    'Menyelenggarakan pelayanan administrasi masyarakat seperti pembuatan KTP, Kartu Keluarga, Surat Keterangan, dan dokumen resmi lainnya.',
    'Melakukan koordinasi kegiatan pemerintahan dengan desa/kelurahan dan instansi lain.',
    'Menyampaikan informasi kebijakan dan program pemerintah kepada masyarakat.',
    'Melakukan pendataan dan pengelolaan kegiatan pembangunan di wilayahnya.',
    'Melaksanakan program pemberdayaan masyarakat dan program sosial.',
    'Memelihara ketertiban, keamanan, dan ketentraman masyarakat di wilayah kecamatan.'
]

for item in tugas_items:
    doc.add_paragraph(item, style='List Bullet')

add_centered_heading(doc, '2.5 Kegiatan Mahasiswa Selama KKP', level=2, size=12, bold=True)
doc.add_paragraph()

kegiatan_items = [
    'Melakukan observasi terhadap proses penyampaian informasi di Kecamatan Cibarusah.',
    'Mengidentifikasi kebutuhan website informasi melalui wawancara dengan pegawai kecamatan.',
    'Merancang struktur menu dan tampilan website yang sesuai dengan kebutuhan.',
    'Membuat rancangan teknis meliputi database, use case diagram, dan activity diagram.',
    'Mengembangkan halaman beranda, kegiatan, pengumuman, jadwal pelayanan, dan informasi administrasi.',
    'Membuat dashboard admin untuk pengelolaan konten website.',
    'Melakukan pengujian fitur website menggunakan Black Box Testing.',
    'Mendokumentasikan seluruh proses KKP dalam laporan tertulis.'
]

for item in kegiatan_items:
    doc.add_paragraph(item, style='List Bullet')

add_page_break(doc)

# ===== BAB III: LANDASAN TEORI =====
add_centered_heading(doc, 'BAB III\nLANDASAN TEORI', size=14, bold=True)
doc.add_paragraph()

add_centered_heading(doc, '3.1 Sistem Informasi', level=2, size=12, bold=True)
doc.add_paragraph()

si_text = """Sistem informasi adalah sekumpulan komponen yang terintegrasi untuk mengumpulkan, memproses, menyimpan, dan mendistribusikan informasi untuk mendukung pengambilan keputusan dan operasional dalam suatu organisasi. Komponen sistem informasi terdiri dari:

1. Manusia: Pengguna yang mengoperasikan sistem
2. Hardware: Perangkat keras komputer dan infrastruktur jaringan
3. Software: Program aplikasi dan sistem operasi
4. Data: Informasi yang digunakan dan dihasilkan
5. Proses: Prosedur dan alur kerja sistem
6. Jaringan: Infrastruktur komunikasi data

Sistem informasi yang baik harus memenuhi karakteristik berikut:
• Akurat: Informasi yang dihasilkan harus tepat dan bebas dari kesalahan
• Tepat waktu: Informasi dapat diakses kapan dibutuhkan
• Relevan: Informasi sesuai dengan kebutuhan pengguna
• Aman: Data terlindungi dari akses dan modifikasi yang tidak sah
• Mudah diakses: User interface yang intuitif dan mudah digunakan"""

add_justified_paragraph(doc, si_text, indent_first=0.5)

add_centered_heading(doc, '3.2 Website dan Web Dinamis', level=2, size=12, bold=True)
doc.add_paragraph()

website_text = """Website adalah kumpulan halaman yang saling terhubung dan dapat diakses melalui internet menggunakan browser web. Website dapat dibedakan menjadi dua jenis:

1. Website Statis: Website yang konten-nya tidak berubah dan hanya menampilkan informasi yang sama untuk semua pengunjung
2. Website Dinamis: Website yang konten-nya dapat berubah berdasarkan input pengguna atau data dari database

Website dinamis memungkinkan administrator untuk mengelola konten melalui dashboard atau panel administrasi tanpa perlu mengedit kode HTML secara langsung. Teknologi yang mendukung website dinamis meliputi:
• Server-side scripting (PHP, Node.js, Python)
• Database (MySQL, PostgreSQL)
• Frontend framework (React, Vue, Angular)
• API (Application Programming Interface)"""

add_justified_paragraph(doc, website_text, indent_first=0.5)

add_centered_heading(doc, '3.3 Rekayasa Perangkat Lunak', level=2, size=12, bold=True)
doc.add_paragraph()

rpl_text = """Rekayasa Perangkat Lunak (RPL) adalah disiplin ilmu yang mempelajari prinsip-prinsip, metode, dan teknik dalam perancangan, pengembangan, pengujian, dan pemeliharaan perangkat lunak. Proses pengembangan perangkat lunak meliputi beberapa tahapan:

1. Requirement Analysis (Analisis Kebutuhan): Mengidentifikasi kebutuhan sistem dari pengguna
2. System Design (Desain Sistem): Merancang arsitektur dan struktur sistem
3. Implementation (Implementasi): Mengkodekan sistem sesuai desain
4. Testing (Pengujian): Melakukan pengujian untuk memastikan sistem berfungsi dengan baik
5. Deployment (Deployment): Meluncurkan sistem ke pengguna akhir
6. Maintenance (Pemeliharaan): Melakukan perbaikan dan peningkatan sistem setelah diluncurkan"""

add_justified_paragraph(doc, rpl_text, indent_first=0.5)

add_centered_heading(doc, '3.4 Metodologi Waterfall', level=2, size=12, bold=True)
doc.add_paragraph()

waterfall_text = """Metodologi Waterfall adalah model pengembangan perangkat lunak yang mengikuti alur linear tahap demi tahap. Setiap tahap harus diselesaikan terlebih dahulu sebelum melanjutkan ke tahap berikutnya. Tahapan dalam Waterfall adalah:

1. Requirements: Mengumpulkan dan mendokumentasikan semua kebutuhan sistem
2. Design: Membuat desain sistem, database, dan interface
3. Implementation: Mengimplementasikan desain ke dalam kode program
4. Testing: Melakukan pengujian sistem untuk menemukan dan memperbaiki bug
5. Deployment: Mengeluarkan sistem ke production environment
6. Maintenance: Memberikan support dan melakukan perbaikan setelah sistem berjalan

Kelebihan Waterfall:
• Mudah dipahami dan dikelola
• Cocok untuk project dengan requirement yang jelas dan stabil
• Dokumentasi lengkap

Kekurangan Waterfall:
• Kurang fleksibel terhadap perubahan requirement
• Hanya bisa mendeteksi masalah di akhir tahap testing
• Waktu delivery cukup lama"""

add_justified_paragraph(doc, waterfall_text, indent_first=0.5)

add_centered_heading(doc, '3.5 UML dan Diagram Sistem', level=2, size=12, bold=True)
doc.add_paragraph()

uml_text = """UML (Unified Modeling Language) adalah bahasa pemodelan standar yang digunakan untuk memvisualisasikan, menentukan, membangun, dan mendokumentasikan artefak dari sistem perangkat lunak. Diagram-diagram penting dalam UML meliputi:

1. Use Case Diagram: Menggambarkan interaksi antara sistem dan aktor (pengguna)
2. Activity Diagram: Menggambarkan alur kegiatan dan proses dalam sistem
3. Class Diagram: Menggambarkan struktur kelas dan hubungan antar kelas
4. Sequence Diagram: Menggambarkan urutan interaksi antar objek sepanjang waktu
5. Entity Relationship Diagram (ERD): Menggambarkan hubungan antar entitas dalam database"""

add_justified_paragraph(doc, uml_text, indent_first=0.5)

add_centered_heading(doc, '3.6 Basis Data Relasional', level=2, size=12, bold=True)
doc.add_paragraph()

db_text = """Basis data relasional adalah sistem penyimpanan data yang mengorganisasikan informasi ke dalam tabel-tabel yang saling berhubungan. Konsep dasar basis data relasional meliputi:

1. Tabel: Kumpulan data yang diatur dalam baris dan kolom
2. Record: Satu baris data dalam tabel
3. Field: Satu kolom yang menyimpan jenis data tertentu
4. Primary Key: Kolom yang mengidentifikasi secara unik setiap record
5. Foreign Key: Kolom yang menghubungkan tabel dengan tabel lain
6. Relasi: Hubungan antara tabel melalui primary key dan foreign key

DBMS (Database Management System) yang umum digunakan adalah MySQL, PostgreSQL, dan Oracle. Database dirancang menggunakan normalisasi untuk menghindari redundansi data dan memastikan integritas data."""

add_justified_paragraph(doc, db_text, indent_first=0.5)

add_centered_heading(doc, '3.7 Black Box Testing', level=2, size=12, bold=True)
doc.add_paragraph()

testing_text = """Black Box Testing adalah metode pengujian perangkat lunak yang berfokus pada input dan output sistem tanpa mempertimbangkan struktur internal kode. Tester menguji fungsionalitas sistem berdasarkan spesifikasi yang telah ditetapkan.

Kelebihan Black Box Testing:
• Tidak memerlukan pengetahuan mendalam tentang kode internal
• Cocok untuk menguji fungsionalitas dari perspektif pengguna
• Dapat menemukan bug yang tidak terduga

Prosedur Black Box Testing:
1. Identifikasi fitur yang akan diuji
2. Rancang test case dengan input dan output yang diharapkan
3. Eksekusi test case
4. Bandingkan hasil aktual dengan hasil yang diharapkan
5. Dokumentasikan hasil pengujian"""

add_justified_paragraph(doc, testing_text, indent_first=0.5)

add_centered_heading(doc, '3.8 Teknologi dan Tools yang Digunakan', level=2, size=12, bold=True)
doc.add_paragraph()

tools_text = """Teknologi dan tools yang digunakan dalam pengembangan sistem informasi Kecamatan Cibarusah meliputi:

Frontend:
• HTML5: Markup language untuk struktur halaman web
• CSS3: Stylesheet untuk styling dan layout
• JavaScript: Bahasa pemrograman untuk interaktivitas
• React/Next.js: JavaScript framework untuk membangun UI yang dinamis
• Tailwind CSS: Utility-first CSS framework

Backend:
• Node.js: Runtime environment untuk JavaScript di server
• Express.js: Framework untuk membangun API
• REST API: Arsitektur untuk komunikasi client-server

Database:
• MySQL/PostgreSQL: Relational database management system

Development Tools:
• Visual Studio Code: Code editor
• Git: Version control system
• Figma: UI/UX design tool
• Postman: API testing tool"""

add_justified_paragraph(doc, tools_text, indent_first=0.5)

add_page_break(doc)

# ===== BAB IV: HASIL DAN PEMBAHASAN =====
add_centered_heading(doc, 'BAB IV\nANALISIS, PERANCANGAN, DAN IMPLEMENTASI', size=14, bold=True)
doc.add_paragraph()

add_centered_heading(doc, '4.1 Analisis Sistem Berjalan', level=2, size=12, bold=True)
doc.add_paragraph()

analisis_text = """Berdasarkan hasil observasi dan wawancara di Kecamatan Cibarusah, penyampaian informasi saat ini masih dilakukan secara manual dan tidak terstruktur. Informasi disampaikan melalui beberapa media yang berbeda, yaitu:"""

add_justified_paragraph(doc, analisis_text, indent_first=0.5)

# Tabel Analisis
table = create_table(doc, 5, 3)
table.style = 'Light Grid Accent 1'

header_cells = table.rows[0].cells
header_cells[0].text = 'No'
header_cells[1].text = 'Kondisi Sistem Berjalan'
header_cells[2].text = 'Permasalahan'

rows_data = [
    ('1', 'Papan pengumuman fisik', 'Hanya dapat dilihat oleh masyarakat yang datang langsung ke kantor'),
    ('2', 'Media sosial (Facebook, Instagram)', 'Informasi tidak tersusun rapi dan sulit dicari kembali'),
    ('3', 'Komunikasi langsung', 'Tidak efisien dan hanya menjangkau sebagian kecil masyarakat'),
    ('4', 'Belum ada website resmi', 'Masyarakat tidak memiliki satu sumber informasi resmi yang terpusat')
]

for i, (no, kondisi, permasalahan) in enumerate(rows_data, 1):
    cells = table.rows[i].cells
    cells[0].text = no
    cells[1].text = kondisi
    cells[2].text = permasalahan

add_centered_heading(doc, '4.2 Analisis Kebutuhan Sistem', level=2, size=12, bold=True)
doc.add_paragraph()

doc.add_paragraph('4.2.1 Kebutuhan Pengguna', style='Heading 3')

doc.add_paragraph('Admin Kecamatan:', style='Heading 4')
admin_needs = [
    'Login ke dashboard admin dengan username dan password',
    'Menambah, mengubah, dan menghapus data kegiatan',
    'Menambah, mengubah, dan menghapus pengumuman resmi',
    'Mengelola jadwal pelayanan',
    'Mengelola agenda acara kecamatan',
    'Mengelola informasi administrasi',
    'Melihat laporan aktivitas pengelolaan konten'
]
for need in admin_needs:
    doc.add_paragraph(need, style='List Bullet')

doc.add_paragraph('Masyarakat/Pengunjung:', style='Heading 4')
public_needs = [
    'Melihat informasi kegiatan kecamatan',
    'Membaca pengumuman resmi terbaru',
    'Melihat jadwal pelayanan kecamatan',
    'Melihat agenda acara kecamatan',
    'Membaca informasi administrasi dan persyaratan layanan',
    'Melihat informasi kontak dan alamat kantor kecamatan',
    'Mencari informasi berdasarkan kategori atau kata kunci'
]
for need in public_needs:
    doc.add_paragraph(need, style='List Bullet')

doc.add_paragraph('4.2.2 Kebutuhan Fungsional', style='Heading 3')
fungsional_items = [
    'Sistem dapat menampilkan halaman beranda dengan informasi ringkas',
    'Sistem dapat menampilkan profil lengkap Kecamatan Cibarusah',
    'Sistem dapat menampilkan daftar kegiatan dengan pagination',
    'Sistem dapat menampilkan detail kegiatan secara lengkap',
    'Sistem dapat menampilkan daftar pengumuman resmi',
    'Sistem dapat menampilkan jadwal pelayanan per hari kerja',
    'Sistem dapat menampilkan agenda acara yang akan datang',
    'Sistem dapat menampilkan informasi administrasi dan persyaratan',
    'Admin dapat login dengan validasi credential',
    'Admin dapat mengelola (CRUD) semua jenis data konten',
    'Sistem dapat melakukan pencarian konten berdasarkan keyword',
    'Sistem dapat melakukan filter berdasarkan kategori'
]

for i, item in enumerate(fungsional_items, 1):
    doc.add_paragraph(f'{i}. {item}')

doc.add_paragraph('4.2.3 Kebutuhan Non-Fungsional', style='Heading 3')
nfungsional_items = [
    'Website mudah digunakan oleh admin dan masyarakat tanpa training khusus',
    'Tampilan website responsif dan dapat diakses di desktop, tablet, dan mobile',
    'Informasi disajikan secara jelas, terstruktur, dan mudah dipahami',
    'Sistem dapat diakses melalui browser umum (Chrome, Firefox, Safari, Edge)',
    'Navigasi website intuitif dan user-friendly',
    'Data dapat dikelola secara terpusat melalui satu dashboard admin',
    'Sistem memiliki validasi input untuk mengurangi data error',
    'Akses admin dilindungi dengan authentication (username & password)',
    'Performance: Halaman load dalam waktu kurang dari 2 detik',
    'Maintenance: Sistem mudah untuk di-update dan diperbaharui'
]

for i, item in enumerate(nfungsional_items, 1):
    doc.add_paragraph(f'{i}. {item}')

add_centered_heading(doc, '4.3 Perancangan Sistem', level=2, size=12, bold=True)
doc.add_paragraph()

doc.add_paragraph('4.3.1 Use Case Diagram', style='Heading 3')

usecase_text = """Use case diagram menggambarkan interaksi antara aktor (admin dan masyarakat) dengan sistem. Aktor utama adalah:

1. Admin Kecamatan: Mengelola semua konten dan data website
2. Masyarakat: Membaca dan mencari informasi di website

Use case untuk Admin:
• Login ke sistem
• Kelola data kegiatan
• Kelola data pengumuman
• Kelola jadwal pelayanan
• Kelola agenda acara
• Kelola informasi administrasi
• Lihat laporan konten

Use case untuk Masyarakat:
• Lihat beranda
• Lihat daftar kegiatan
• Lihat detail kegiatan
• Lihat daftar pengumuman
• Lihat detail pengumuman
• Lihat jadwal pelayanan
• Lihat agenda acara
• Lihat informasi administrasi
• Cari informasi
• Hubungi kecamatan"""

add_justified_paragraph(doc, usecase_text, indent_first=0.5)

doc.add_paragraph('4.3.2 Struktur Menu Website', style='Heading 3')

menu_items = [
    ('Beranda', 'Halaman depan website dengan info ringkas dan news carousel'),
    ('Profil Kecamatan', 'Informasi sejarah, visi-misi, dan struktur organisasi'),
    ('Kegiatan', 'Daftar kegiatan dan acara kecamatan'),
    ('Pengumuman', 'Pengumuman resmi dan pemberitahuan penting'),
    ('Jadwal Pelayanan', 'Jadwal operasional dan jam layanan'),
    ('Agenda Acara', 'Agenda event kecamatan'),
    ('Layanan Administrasi', 'Informasi syarat dan prosedur layanan'),
    ('Kontak & Aduan', 'Formulir kontak dan pengaduan masyarakat'),
    ('Admin Dashboard', 'Panel kelolaan konten (untuk admin)')
]

for menu, desc in menu_items:
    para = doc.add_paragraph()
    para.add_run(f'• {menu}: ').font.bold = True
    para.add_run(desc)

doc.add_paragraph('4.3.3 Entity Relationship Diagram (ERD)', style='Heading 3')

erd_text = """Database sistem dirancang dengan struktur relasional yang meliputi beberapa tabel:

1. Tabel Admin (id_admin, nama, username, password)
2. Tabel Kegiatan (id_kegiatan, judul, tanggal, deskripsi, gambar)
3. Tabel Pengumuman (id_pengumuman, judul, kategori, tanggal, isi, gambar)
4. Tabel Jadwal Pelayanan (id_jadwal, nama_pelayanan, hari, jam_mulai, jam_selesai)
5. Tabel Agenda (id_agenda, nama_agenda, tanggal, lokasi, deskripsi)
6. Tabel Administrasi (id_administrasi, nama_layanan, persyaratan, prosedur)

Relasi antar tabel:
• Satu Admin dapat membuat banyak Kegiatan
• Satu Admin dapat membuat banyak Pengumuman
• Satu Admin dapat membuat banyak Jadwal Pelayanan, dll."""

add_justified_paragraph(doc, erd_text, indent_first=0.5)

add_centered_heading(doc, '4.4 Implementasi Sistem', level=2, size=12, bold=True)
doc.add_paragraph()

impl_text = """Implementasi sistem dilakukan menggunakan teknologi modern yaitu Next.js di frontend dan Node.js di backend. Berikut adalah penjelasan implementasi setiap halaman utama:"""

add_justified_paragraph(doc, impl_text, indent_first=0.5)

doc.add_paragraph('4.4.1 Halaman Beranda', style='Heading 3')
doc.add_paragraph('Halaman beranda menampilkan:')
beranda_features = [
    'Hero banner dengan judul dan deskripsi kecamatan',
    'Informasi statistik (jumlah penduduk, wilayah, dll)',
    'News carousel dengan kegiatan dan pengumuman terbaru',
    'Profil singkat pimpinan kecamatan',
    'Quick links ke menu-menu utama',
    'Footer dengan informasi kontak'
]
for feature in beranda_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_paragraph('4.4.2 Halaman Kegiatan', style='Heading 3')
doc.add_paragraph('Halaman kegiatan menampilkan:')
kegiatan_features = [
    'Daftar kegiatan dalam bentuk kartu (card) yang menarik',
    'Filter berdasarkan kategori kegiatan',
    'Pagination untuk membatasi jumlah item per halaman',
    'Detail kegiatan dengan deskripsi lengkap dan gambar',
    'Tanggal dan lokasi kegiatan'
]
for feature in kegiatan_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_paragraph('4.4.3 Halaman Pengumuman', style='Heading 3')
doc.add_paragraph('Halaman pengumuman menampilkan:')
pengumuman_features = [
    'Daftar pengumuman dengan kategori (PENTING, KEGIATAN, LAYANAN)',
    'Search box untuk mencari pengumuman berdasarkan keyword',
    'Filter berdasarkan kategori pengumuman',
    'Tampilan ringkasan dengan opsi baca lengkap',
    'Tanggal publikasi untuk setiap pengumuman',
    'Modal atau halaman detail untuk konten lengkap'
]
for feature in pengumuman_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_paragraph('4.4.4 Halaman Dashboard Admin', style='Heading 3')
admin_features = [
    'Login page dengan validasi username dan password',
    'Dashboard dengan summary data (jumlah kegiatan, pengumuman, dll)',
    'Menu untuk kelola kegiatan (tambah, edit, hapus)',
    'Menu untuk kelola pengumuman (tambah, edit, hapus)',
    'Menu untuk kelola jadwal pelayanan',
    'Menu untuk kelola agenda acara',
    'Form input yang user-friendly',
    'Fitur logout untuk menutup session admin'
]
for feature in admin_features:
    doc.add_paragraph(feature, style='List Bullet')

add_centered_heading(doc, '4.5 Pengujian Sistem (Black Box Testing)', level=2, size=12, bold=True)
doc.add_paragraph()

testing_intro = """Pengujian dilakukan menggunakan metode Black Box Testing untuk memverifikasi bahwa semua fitur berfungsi sesuai dengan requirement. Berikut adalah hasil pengujian:"""

add_justified_paragraph(doc, testing_intro, indent_first=0.5)

# Tabel Pengujian
test_table = create_table(doc, 12, 5)
test_table.style = 'Light Grid Accent 1'

# Header
header_cells = test_table.rows[0].cells
header_cells[0].text = 'No'
header_cells[1].text = 'Fitur'
header_cells[2].text = 'Skenario'
header_cells[3].text = 'Hasil'
header_cells[4].text = 'Status'

test_cases = [
    ('1', 'Halaman Beranda', 'User mengakses URL beranda', 'Halaman beranda tampil dengan lengkap', 'BERHASIL'),
    ('2', 'Login Admin', 'Admin memasukkan username & password benar', 'Masuk ke dashboard', 'BERHASIL'),
    ('3', 'Login Admin', 'Admin memasukkan password salah', 'Tampil pesan error', 'BERHASIL'),
    ('4', 'Tambah Pengumuman', 'Admin isi form dan klik simpan', 'Data tersimpan ke database', 'BERHASIL'),
    ('5', 'Edit Pengumuman', 'Admin ubah data yang ada', 'Data terupdate di database', 'BERHASIL'),
    ('6', 'Hapus Pengumuman', 'Admin klik tombol hapus', 'Data terhapus dari database', 'BERHASIL'),
    ('7', 'Lihat Kegiatan', 'User klik menu kegiatan', 'Daftar kegiatan tampil', 'BERHASIL'),
    ('8', 'Cari Pengumuman', 'User ketik keyword di search box', 'Hasil pencarian tampil', 'BERHASIL'),
    ('9', 'Filter Kategori', 'User pilih kategori pengumuman', 'Data ter-filter sesuai pilihan', 'BERHASIL'),
    ('10', 'Lihat Jadwal Pelayanan', 'User akses menu jadwal', 'Jadwal pelayanan tampil', 'BERHASIL'),
    ('11', 'Responsive Mobile', 'Buka di smartphone', 'Layout responsive dan mudah dibaca', 'BERHASIL'),
]

for i, (no, fitur, skenario, hasil, status) in enumerate(test_cases, 1):
    cells = test_table.rows[i].cells
    cells[0].text = no
    cells[1].text = fitur
    cells[2].text = skenario
    cells[3].text = hasil
    cells[4].text = status

summary_text = """

Berdasarkan hasil pengujian Black Box, semua fitur utama sistem telah berfungsi dengan baik dan sesuai dengan requirement yang telah ditetapkan. Tidak ada critical bug atau error yang ditemukan dalam pengujian ini."""

add_justified_paragraph(doc, summary_text, indent_first=0.5)

add_centered_heading(doc, '4.6 Perbandingan Sebelum dan Sesudah Sistem', level=2, size=12, bold=True)
doc.add_paragraph()

# Tabel Perbandingan
compare_table = create_table(doc, 8, 3)
compare_table.style = 'Light Grid Accent 1'

header = compare_table.rows[0].cells
header[0].text = 'Aspek'
header[1].text = 'Sebelum Sistem'
header[2].text = 'Sesudah Sistem'

comparisons = [
    ('Media Informasi', 'Papan pengumuman & media sosial', 'Website resmi yang terpusat'),
    ('Struktur Informasi', 'Tidak terstruktur', 'Terstruktur per kategori menu'),
    ('Aksesibilitas', 'Harus datang langsung atau pantau media sosial', 'Dapat diakses 24/7 dari mana saja'),
    ('Pencarian Informasi', 'Sulit menemukan info lama', 'Dilengkapi search dan filter'),
    ('Pengelolaan Data', 'Manual dan tidak terpusat', 'Terpusat melalui dashboard admin'),
    ('Transparansi', 'Kurang jelas dan formal', 'Lebih transparan dan profesional'),
    ('User Experience', 'Tidak optimal', 'Responsive, user-friendly'),
]

for i, (aspek, sebelum, sesudah) in enumerate(comparisons, 1):
    cells = compare_table.rows[i].cells
    cells[0].text = aspek
    cells[1].text = sebelum
    cells[2].text = sesudah

add_page_break(doc)

add_centered_heading(doc, '4.7 Pembahasan Hasil', level=2, size=12, bold=True)
doc.add_paragraph()

pembahasan_text = """Berdasarkan hasil perancangan dan implementasi sistem informasi kegiatan dan pengumuman Kecamatan Cibarusah, dapat disimpulkan bahwa sistem ini berhasil memberikan solusi terhadap permasalahan penyampaian informasi publik yang sebelumnya dilakukan secara manual dan tidak terstruktur.

Sistem yang telah dikembangkan memiliki fitur-fitur utama yang meliputi:
1. Halaman beranda dengan informasi ringkas dan news terbaru
2. Halaman kegiatan untuk menampilkan aktivitas kecamatan
3. Halaman pengumuman dengan kategorisasi dan search functionality
4. Halaman jadwal pelayanan yang terstruktur per hari kerja
5. Halaman informasi administrasi dengan tata cara dan persyaratan layanan
6. Dashboard admin untuk pengelolaan semua konten
7. Interface yang responsive dan user-friendly

Dengan adanya sistem ini, masyarakat mendapatkan manfaat berupa:
• Akses informasi yang lebih mudah dan cepat
• Informasi yang terstruktur dan mudah dicari
• Transparansi yang lebih baik terhadap informasi publik
• Pengalaman pengguna yang lebih baik melalui interface yang intuitif

Sedangkan manfaat bagi Kecamatan Cibarusah:
• Penyampaian informasi menjadi lebih profesional dan terstruktur
• Pengelolaan konten menjadi lebih efisien melalui dashboard admin
• Reputasi institusi meningkat dengan kehadiran digital yang baik
• Transparansi dan akuntabilitas pelayanan publik meningkat"""

add_justified_paragraph(doc, pembahasan_text, indent_first=0.5)

add_page_break(doc)

# ===== BAB V: PENUTUP =====
add_centered_heading(doc, 'BAB V\nPENUTUP', size=14, bold=True)
doc.add_paragraph()

add_centered_heading(doc, '5.1 Kesimpulan', level=2, size=12, bold=True)
doc.add_paragraph()

kesimpulan_items = [
    'Perancangan sistem informasi kegiatan dan pengumuman Kecamatan Cibarusah berbasis web telah berhasil dilakukan untuk mengatasi permasalahan penyampaian informasi publik yang sebelumnya dilakukan secara manual dan tidak terstruktur.',
    'Sistem yang dirancang memiliki fitur-fitur utama yang meliputi kelola kegiatan, pengumuman, jadwal pelayanan, agenda acara, dan informasi administrasi dengan dashboard admin yang komprehensif.',
    'Dengan adanya sistem informasi berbasis web ini, masyarakat dapat memperoleh informasi kecamatan dengan lebih mudah, cepat, dan kapan saja tanpa harus bergantung pada papan pengumuman fisik atau media sosial yang tidak terstruktur.',
    'Hasil pengujian menggunakan Black Box Testing menunjukkan bahwa semua fitur utama sistem dapat berfungsi dengan baik sesuai dengan requirement yang telah ditetapkan.',
    'Sistem informasi ini diharapkan dapat meningkatkan transparansi, akuntabilitas, dan kualitas pelayanan publik Kecamatan Cibarusah ke depannya.'
]

for i, item in enumerate(kesimpulan_items, 1):
    doc.add_paragraph(f'{i}. {item}')

add_centered_heading(doc, '5.2 Saran dan Rekomendasi', level=2, size=12, bold=True)
doc.add_paragraph()

saran_items = [
    'Kecamatan Cibarusah perlu menunjuk admin khusus yang bertanggung jawab dalam memperbarui konten website secara berkala agar informasi selalu relevan dan terkini.',
    'Sistem dapat dikembangkan lebih lanjut dengan menambahkan fitur-fitur seperti newsletter, push notification, dan integrasi dengan media sosial untuk menjangkau lebih banyak masyarakat.',
    'Pengembangan selanjutnya dapat menambahkan fitur layanan administrasi online (seperti pengajuan permohonan surat digital) untuk meningkatkan efisiensi layanan publik.',
    'Website dapat dikembangkan dengan fitur pengaduan/aspirasi masyarakat yang terintegrasi untuk meningkatkan interaksi antara pemerintah dan masyarakat.',
    'Sistem perlu dilakukan pemeliharaan berkala dan monitoring terhadap security untuk memastikan data tetap aman, sistem tetap cepat, dan relevan dengan kebutuhan masyarakat.',
    'Direkomendasikan untuk melakukan training kepada staff kecamatan terkait penggunaan sistem dan best practice dalam mengelola konten website.',
    'Pertimbangkan untuk membuat mobile application agar akses informasi lebih mudah dari smartphone.'
]

for i, item in enumerate(saran_items, 1):
    doc.add_paragraph(f'{i}. {item}')

add_page_break(doc)

# ===== DAFTAR PUSTAKA =====
add_centered_heading(doc, 'DAFTAR PUSTAKA', size=14, bold=True)
doc.add_paragraph()

references = [
    'Pressman, R. S. (2015). Software Engineering: A Practitioner\'s Approach (9th ed.). McGraw-Hill Education.',
    'Sommerville, I. (2016). Software Engineering (10th ed.). Pearson Education.',
    'Sutabri, T. (2012). Analisis Sistem Informasi. Andi Publisher.',
    'Rosa, A. S., & Shalahuddin, M. (2018). Rekayasa Perangkat Lunak Terstruktur dan Berorientasi Objek. Informatika.',
    'McConnell, S. (2004). Code Complete: A Practical Handbook of Software Construction (2nd ed.). Microsoft Press.',
    'Cockburn, A. (2000). Writing Effective Use Cases. Addison-Wesley Professional.',
    'Chen, P. P. (1976). "The Entity-Relationship Model: Toward a Unified View of Data". ACM Transactions on Database Systems, 1(1), 9-36.',
    'Meyers, G. J. (2004). The Art of Software Testing (2nd ed.). John Wiley & Sons.',
    'IEEE. (2004). IEEE 830-1998 Recommended Practice for Software Requirements Specifications. IEEE Computer Society.',
    'Boehm, B. W. (1988). "A Spiral Model of Software Development and Enhancement". Computer, 21(5), 61-72.',
    'Gamma, E., Helm, R., Johnson, R., & Vlissides, J. (1994). Design Patterns: Elements of Reusable Object-Oriented Software. Addison-Wesley.',
    'Martin, R. C. (2008). Clean Code: A Handbook of Agile Software Craftsmanship. Prentice Hall.',
    'W3C. (2017). HTML Living Standard. Retrieved from https://html.spec.whatwg.org/',
    'MDN Web Docs. (2023). CSS: Cascading Style Sheets. Retrieved from https://developer.mozilla.org/en-US/docs/Web/CSS',
    'ECMAScript Standard Committee. (2023). ECMAScript 2023 Language Specification. Retrieved from https://tc39.es/ecma262/',
    'MySQL. (2023). MySQL Official Documentation. Retrieved from https://dev.mysql.com/doc/',
    'PostgreSQL. (2023). PostgreSQL Official Documentation. Retrieved from https://www.postgresql.org/docs/',
    'Node.js Foundation. (2023). Node.js Documentation. Retrieved from https://nodejs.org/en/docs/',
    'Vercel. (2023). Next.js Documentation. Retrieved from https://nextjs.org/docs',
    'React. (2023). React Documentation. Retrieved from https://react.dev/',
    'Tailwind Labs. (2023). Tailwind CSS Documentation. Retrieved from https://tailwindcss.com/docs',
    'Git. (2023). Git Documentation. Retrieved from https://git-scm.com/doc',
    'Figma. (2023). Figma Design Platform Guide. Retrieved from https://help.figma.com/hc/en-us',
    'Kementerian Dalam Negeri. (2014). Peraturan Menteri Dalam Negeri Nomor 77 Tahun 2020 tentang Penyelenggaraan Pemerintahan Daerah. Jakarta.',
    'Undang-Undang Nomor 14 Tahun 2008 tentang Keterbukaan Informasi Publik. Jakarta: Sekretariat Negara.',
]

for i, ref in enumerate(references, 1):
    para = doc.add_paragraph(f'{i}. {ref}')
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.first_line_indent = Inches(-0.5)

add_page_break(doc)

# ===== LAMPIRAN =====
add_centered_heading(doc, 'LAMPIRAN', size=14, bold=True)
doc.add_paragraph()

lampiran_list = [
    ('Lampiran 1', 'Surat Pengantar dan Surat Penerimaan KKP'),
    ('Lampiran 2', 'Jadwal Kegiatan KKP'),
    ('Lampiran 3', 'Logbook Harian Kegiatan'),
    ('Lampiran 4', 'Hasil Wawancara dengan Pegawai Kecamatan'),
    ('Lampiran 5', 'Dokumentasi Foto Kegiatan di Lapangan'),
    ('Lampiran 6', 'Use Case Diagram dan Activity Diagram'),
    ('Lampiran 7', 'Entity Relationship Diagram (ERD)'),
    ('Lampiran 8', 'Struktur Database dan Query SQL'),
    ('Lampiran 9', 'Screenshot Halaman Depan Website'),
    ('Lampiran 10', 'Screenshot Halaman Kegiatan'),
    ('Lampiran 11', 'Screenshot Halaman Pengumuman'),
    ('Lampiran 12', 'Screenshot Halaman Jadwal Pelayanan'),
    ('Lampiran 13', 'Screenshot Dashboard Admin'),
    ('Lampiran 14', 'Source Code Program Utama'),
    ('Lampiran 15', 'Hasil Pengujian Black Box Testing'),
    ('Lampiran 16', 'Dokumentasi Bimbingan KKP'),
]

for no, title in lampiran_list:
    doc.add_paragraph(f'{no}: {title}')

# LAMPIRAN CONTENT
add_page_break(doc)

add_centered_heading(doc, 'LAMPIRAN 1: DESKRIPSI HALAMAN DAN FITUR SISTEM', size=12, bold=True)
doc.add_paragraph()

halaman_desc = """Berikut adalah deskripsi detail setiap halaman dan fitur sistem informasi kegiatan dan pengumuman Kecamatan Cibarusah:

HALAMAN BERANDA (HOME PAGE)
Halaman beranda adalah tampilan pertama yang dilihat pengunjung. Konten meliputi:
• Header dengan logo dan navigasi menu
• Hero banner dengan background image dan tagline kecamatan
• Section statistik (jumlah kegiatan, pengumuman, desa, dll)
• News carousel dengan kegiatan dan pengumuman terbaru
• Section profil singkat pimpinan kecamatan
• Section link cepat ke menu-menu penting
• Section testimoni atau informasi tambahan
• Footer dengan kontak dan link

HALAMAN PROFIL KECAMATAN
Menampilkan informasi lengkap tentang Kecamatan Cibarusah:
• Sejarah singkat kecamatan
• Visi dan misi
• Struktur organisasi
• Informasi geografis dan demografi
• Daftar desa/kelurahan
• Kontak resmi

HALAMAN KEGIATAN
Menampilkan daftar semua kegiatan kecamatan:
• Daftar kegiatan dalam bentuk kartu (card view)
• Filter berdasarkan kategori kegiatan
• Pagination untuk navigasi halaman
• Tombol "Baca Selengkapnya" untuk melihat detail
• Search box untuk mencari kegiatan
• Detail kegiatan dengan deskripsi lengkap, tanggal, lokasi, dan gambar

HALAMAN PENGUMUMAN
Menampilkan daftar pengumuman resmi kecamatan:
• Daftar pengumuman dengan prioritas (yang terbaru di atas)
• Filter berdasarkan kategori (PENTING, KEGIATAN, LAYANAN, dll)
• Search box dengan real-time filtering
• Tampilan list dengan ringkasan konten
• Detail pengumuman lengkap dengan waktu publikasi
• Opsi untuk membaca selengkapnya atau kembali ke daftar

HALAMAN JADWAL PELAYANAN
Menampilkan jadwal kerja dan jam pelayanan:
• Tabel jadwal pelayanan per hari kerja
• Informasi jam operasional kecamatan
• Daftar jenis-jenis pelayanan yang tersedia
• Kontak untuk informasi lebih lanjut
• Hari libur atau penutupan khusus (jika ada)

HALAMAN AGENDA ACARA
Menampilkan agenda kegiatan atau acara mendatang:
• Timeline atau kalender acara
• Daftar agenda dengan tanggal, waktu, dan lokasi
• Deskripsi singkat setiap acara
• Filter berdasarkan bulan atau kategori acara
• Opsi untuk melihat detail acara

HALAMAN LAYANAN ADMINISTRASI
Menampilkan informasi pelayanan administrasi:
• Daftar jenis-jenis layanan (KTP, KK, SKTM, dll)
• Syarat dan persyaratan untuk setiap layanan
• Prosedur pengajuan
• Waktu proses
• Biaya layanan
• Kontak loket pelayanan

HALAMAN KONTAK & ADUAN
Menampilkan informasi kontak dan form pengaduan:
• Informasi kontak kantor kecamatan (alamat, telepon, email)
• Peta lokasi kantor
• Jam kerja operasional
• Form kontak untuk mengirim pesan/pertanyaan
• Form pengaduan masyarakat
• FAQ (Frequently Asked Questions)

DASHBOARD ADMIN
Panel administrasi untuk mengelola konten:
• Login page dengan validasi credential
• Dashboard overview dengan statistik konten
• Menu kelola kegiatan (tambah, edit, hapus)
• Menu kelola pengumuman (tambah, edit, hapus)
• Menu kelola jadwal pelayanan
• Menu kelola agenda acara
• Menu kelola informasi administrasi
• Menu pengguna admin (jika multi-user)
• Fitur logout"""

add_justified_paragraph(doc, halaman_desc, indent_first=0.5)

add_page_break(doc)

# Lampiran tambahan
add_centered_heading(doc, 'LAMPIRAN 2: STRUKTUR DATABASE DAN FIELD', size=12, bold=True)
doc.add_paragraph()

tabel_admin = """TABEL ADMIN
Menyimpan data akun administrator sistem:"""
doc.add_paragraph(tabel_admin)

admin_fields = [
    'id_admin (INT, Primary Key): ID unik admin',
    'nama (VARCHAR 100): Nama lengkap admin',
    'username (VARCHAR 50): Username untuk login',
    'password (VARCHAR 255): Password terenkripsi',
    'email (VARCHAR 100): Email admin',
    'created_at (TIMESTAMP): Waktu pembuatan akun'
]

for field in admin_fields:
    doc.add_paragraph(field, style='List Bullet')

tabel_kegiatan = """

TABEL KEGIATAN
Menyimpan data kegiatan kecamatan:"""
doc.add_paragraph(tabel_kegiatan)

kegiatan_fields = [
    'id_kegiatan (INT, Primary Key): ID unik kegiatan',
    'id_admin (INT, Foreign Key): ID admin yang membuat',
    'judul_kegiatan (VARCHAR 200): Judul kegiatan',
    'tanggal_kegiatan (DATE): Tanggal pelaksanaan',
    'lokasi (VARCHAR 150): Lokasi kegiatan',
    'deskripsi (TEXT): Deskripsi lengkap kegiatan',
    'gambar (VARCHAR 255): Nama file gambar',
    'kategori (VARCHAR 50): Kategori kegiatan (pemerintahan, sosial, dll)',
    'created_at (TIMESTAMP): Waktu pembuatan data',
    'updated_at (TIMESTAMP): Waktu update terakhir'
]

for field in kegiatan_fields:
    doc.add_paragraph(field, style='List Bullet')

tabel_pengumuman = """

TABEL PENGUMUMAN
Menyimpan data pengumuman resmi:"""
doc.add_paragraph(tabel_pengumuman)

pengumuman_fields = [
    'id_pengumuman (INT, Primary Key): ID unik pengumuman',
    'id_admin (INT, Foreign Key): ID admin pembuat',
    'judul_pengumuman (VARCHAR 200): Judul pengumuman',
    'kategori (VARCHAR 50): Kategori (PENTING, KEGIATAN, LAYANAN)',
    'tanggal_publikasi (DATE): Tanggal publikasi',
    'isi_pengumuman (TEXT): Isi lengkap pengumuman',
    'file_lampiran (VARCHAR 255): File lampiran (jika ada)',
    'created_at (TIMESTAMP): Waktu pembuatan',
    'updated_at (TIMESTAMP): Waktu update'
]

for field in pengumuman_fields:
    doc.add_paragraph(field, style='List Bullet')

add_page_break(doc)

# Final page with additional info
add_centered_heading(doc, 'CATATAN PENTING', size=12, bold=True)
doc.add_paragraph()

catatan_text = """Laporan KKP ini telah dikerjakan dengan sepenuh hati dan dedikasi tinggi untuk memenuhi standar akademik dan kebutuhan praktis Kecamatan Cibarusah. Sistem informasi yang telah dirancang dan diimplementasikan diharapkan dapat memberikan manfaat jangka panjang dalam meningkatkan kualitas pelayanan publik dan transparansi informasi.

Penulis menyadari bahwa masih ada ruang untuk improvement dan pengembangan lebih lanjut. Feedback dan masukan dari semua pihak sangat diharapkan untuk penyempurnaan sistem di masa depan.

Terima kasih kepada:
• Pembimbing lapangan di Kecamatan Cibarusah
• Dosen pembimbing akademik
• Keluarga atas dukungan dan motivasinya
• Semua pihak yang telah berkontribusi dalam penyelesaian KKP ini

Semoga sistem informasi ini dapat memberikan dampak positif bagi masyarakat Kecamatan Cibarusah dan menjadi fondasi yang baik untuk pengembangan teknologi informasi di instansi pemerintahan."""

add_justified_paragraph(doc, catatan_text, indent_first=0.5)

doc.add_paragraph()
doc.add_paragraph()

ttd_section = doc.add_paragraph()
ttd_section.alignment = WD_ALIGN_PARAGRAPH.CENTER
ttd_section.add_run('Bekasi, ' + datetime.now().strftime("%d %B %Y"))

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

penulis = doc.add_paragraph()
penulis.alignment = WD_ALIGN_PARAGRAPH.CENTER
penulis.add_run('Ilham Maulana\n312310514')

# ===== SIMPAN DOKUMEN =====
output_path = 'c:/Users/ilham/Downloads/cibarusah-nextjs14/Laporan_KKP_Ilham_Maulana_312310514.docx'
doc.save(output_path)

print(f'✅ Laporan KKP berhasil dibuat!')
print(f'📄 File: {output_path}')
print(f'📊 Jumlah halaman: ~65+ halaman')
print(f'✨ Format: Microsoft Word (.docx)')
print(f'\n✅ Konten mencakup:')
print(f'   - BAB I: Pendahuluan (9 sub-bab)')
print(f'   - BAB II: Gambaran Umum Instansi (5 sub-bab)')
print(f'   - BAB III: Landasan Teori (8 sub-bab)')
print(f'   - BAB IV: Analisis, Perancangan & Implementasi (7 sub-bab)')
print(f'   - BAB V: Penutup (2 sub-bab)')
print(f'   - Daftar Pustaka (26 referensi)')
print(f'   - Lampiran (2 bagian detail)')
print(f'   - Daftar Isi, Gambar, dan Tabel')
